"""Helper: fülle ZAG-Vorlage (docx) mit generierten Inhalten.

Die Haupt-Vorlage ist `templates/vorlagen/05_VA_FaGe_Inhaltsverzeichnis_VORLAGE.docx`.
Sie enthält bereits die komplette VA-Struktur inkl. Anhang mit:
- Zeitplan-Tabelle (23 Zeilen)
- 8-10 Wochenjournal-Tabellen (je 6×2)
- Einverständniserklärung
- Eigenleistungserklärung

Strategie: Öffne die Vorlage, ersetze die "Schreiben Sie hier Ihren Text…" Platzhalter
mit unseren generierten Markdown-Inhalten (nach Section-Mapping), fülle Tabellenzellen
direkt. Output: fertige .docx-Datei, die wie eine echte Schüler-Abgabe aussieht.
"""
from __future__ import annotations
import re
import subprocess
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt

from scripts import config
from scripts.coherence import Universe
from scripts.utils import log


VORLAGE_HAUPT = config.PROJECT_ROOT / "templates" / "vorlagen" / "05_VA_FaGe_Inhaltsverzeichnis_VORLAGE.docx"


def _find_para_idx_with_text(doc: Document, needle: str, after: int = 0) -> int | None:
    """Erster Paragraph ab `after`, dessen Text die needle enthält. -1 → nicht gefunden."""
    for i, p in enumerate(doc.paragraphs[after:], start=after):
        if needle in p.text:
            return i
    return None


def _find_heading_idx(doc: Document, heading_text: str, after: int = 0) -> int | None:
    """Findet Paragraph mit exakt diesem Heading-Text (strip + lower)."""
    needle = heading_text.strip().lower()
    for i, p in enumerate(doc.paragraphs[after:], start=after):
        if p.text.strip().lower() == needle:
            return i
    return None


def _replace_paragraph_text(para: Paragraph, new_text: str) -> None:
    """Ersetzt den Inhalt eines Paragraphen, behält Formatierung des ersten Runs."""
    # Alle Runs bis auf den ersten löschen
    runs = para.runs
    if not runs:
        para.add_run(new_text)
        return
    first_run = runs[0]
    first_run.text = new_text
    for r in runs[1:]:
        r.text = ""


def _insert_paragraph_after(para: Paragraph, text: str, style: str | None = None) -> Paragraph:
    """Fügt nach `para` einen neuen Paragraph ein."""
    new_p = para._parent.add_paragraph(text, style=style)
    # Move the new paragraph right after the source
    para._p.addnext(new_p._p)
    return new_p


def _split_markdown_into_blocks(md: str) -> list[dict]:
    """Parst Markdown in einfache Blöcke: paragraphs + headings + lists + quotes.

    Returns list of {"type": "p|h2|h3|li|quote", "text": "..."}.
    Einfach gehalten — keine Inline-Formatierung.
    """
    blocks: list[dict] = []
    current_para: list[str] = []

    def flush():
        if current_para:
            blocks.append({"type": "p", "text": " ".join(current_para).strip()})
            current_para.clear()

    for line in md.splitlines():
        s = line.rstrip()
        if not s.strip():
            flush()
        elif s.startswith("### "):
            flush()
            blocks.append({"type": "h3", "text": s[4:].strip()})
        elif s.startswith("## "):
            flush()
            blocks.append({"type": "h2", "text": s[3:].strip()})
        elif s.startswith("# "):
            flush()
            blocks.append({"type": "h1", "text": s[2:].strip()})
        elif s.startswith(("- ", "* ", "• ")):
            flush()
            blocks.append({"type": "li", "text": s[2:].strip()})
        elif re.match(r"^\s*\d+\.\s", s):
            flush()
            blocks.append({"type": "li", "text": re.sub(r"^\s*\d+\.\s", "", s)})
        elif s.startswith("> "):
            flush()
            blocks.append({"type": "quote", "text": s[2:].strip()})
        else:
            current_para.append(s.strip())
    flush()
    return blocks


def _replace_placeholder_with_blocks(doc: Document, placeholder_needle: str, blocks: list[dict], after_idx: int = 0) -> int | None:
    """Findet Paragraph mit `placeholder_needle` ab `after_idx`, leert ihn und fügt die Blöcke dort ein.

    Returns: Index des ersten eingefügten Blocks (oder None wenn Placeholder nicht gefunden).
    """
    idx = _find_para_idx_with_text(doc, placeholder_needle, after=after_idx)
    if idx is None:
        return None
    target = doc.paragraphs[idx]
    # Ersten Block in den Target-Paragraph schreiben (ersetzt den Platzhalter-Text)
    if not blocks:
        _replace_paragraph_text(target, "")
        return idx

    first = blocks[0]
    _replace_paragraph_text(target, first["text"])
    # Bei Lists/Quote/Heading: Style anpassen wenn möglich
    try:
        if first["type"] == "li":
            target.style = doc.styles["List Paragraph"]
        elif first["type"] in ("h3",):
            target.style = doc.styles.get_style_by_id("Heading3", 1) if False else target.style
    except Exception:
        pass

    # Restliche Blöcke als neue Paragraphen nach target einfügen
    prev = target
    for b in blocks[1:]:
        style = None
        try:
            if b["type"] == "li":
                style = "List Paragraph"
            elif b["type"] == "h3":
                style = "Heading 3"
            elif b["type"] == "h2":
                style = "Heading 2"
        except Exception:
            style = None
        new_p = _insert_paragraph_after(prev, b["text"], style=style)
        prev = new_p
    return idx


def _extract_section(md: str, section_heading_regex: str, next_heading_regex: str = r"^##\s") -> str:
    """Nimmt aus `md` den Text der Section, deren Heading dem Regex entspricht.

    Stoppt beim nächsten Heading, das `next_heading_regex` matcht.
    """
    lines = md.splitlines()
    section_lines: list[str] = []
    in_section = False
    for line in lines:
        if re.match(section_heading_regex, line, flags=re.IGNORECASE):
            in_section = True
            continue
        if in_section and re.match(next_heading_regex, line):
            break
        if in_section:
            section_lines.append(line)
    return "\n".join(section_lines).strip()


def _fill_table_cell(table, row: int, col: int, text: str) -> None:
    try:
        cell = table.rows[row].cells[col]
        cell.text = text
    except (IndexError, Exception) as e:
        log.warning(f"Tabellenzelle R{row}C{col} konnte nicht gefüllt werden: {e}")


def fill_va_haupt_docx(
    u: Universe,
    haupttext_md: str,
    konzept_md: str,
    journal_md: str,
    zwischenreflexion1_md: str,
    zwischenreflexion2_md: str,
    gesamtreflexion_md: str,
    out_path: Path,
) -> Path:
    """Füllt die Haupt-Vorlage mit allen Inhalten und schreibt .docx.

    Strategie:
    - Sections im Haupttext-Markdown werden den Vorlagen-Absätzen mit "Schreiben Sie hier…"
      zugeordnet und dort eingefügt.
    - Anhang-Tabellen werden explizit gefüllt (Zeitplan, 8 Wochenjournale).
    - Einverständnis- + Eigenleistungs-Felder werden ersetzt.
    """
    doc = Document(str(VORLAGE_HAUPT))

    # === Titel ergänzen (Anfang Para 0) ===
    # Die Vorlage beginnt mit Inhaltsverzeichnis-Struktur. Wir legen den Thema-Titel als ersten Heading davor.
    # Einfacher: ersetze den ersten Heading-Paragraph, falls der generisch ist.

    # === Section-Mapping (Markdown → Vorlagen-Platzhalter) ===
    # Wir suchen die Platzhalter in der Reihenfolge, wie sie im Dokument vorkommen
    mapping = [
        # (Vorlagen-Section-Überschrift oder Context, Markdown-Section-Regex, Placeholder-Needle, "search after")
        {"ctx": "Einleitung",                         "md_re": r"^##\s*1\.?\s*Einleitung", "ph": "Schreiben Sie hier Ihre Einleitung"},
        {"ctx": "2.1 Zahlen und Fakten",              "md_re": r"^###?\s*(2\.1|Zahlen\s+und\s+Fakten)", "ph": "Schreiben Sie hier Ihren Text…", "section_after": "Theoretischer Teil"},
        {"ctx": "2.2 Hintergründe",                   "md_re": r"^###?\s*(2\.2|Hintergründe)", "ph": "Schreiben Sie hier Ihren Text und fügen Sie Bilder ein"},
        {"ctx": "3.1 Erfahrungsbericht",              "md_re": r"^###?\s*(3\.1|Erfahrungsbericht)", "ph": "Schreiben Sie hier Ihren Text…", "section_after": "Erfahrungsbericht"},
        {"ctx": "3.2 Eigene Fotos",                   "md_re": r"^###?\s*(3\.2|Eigene\s+Fotos)", "ph": "Schreiben Sie hier Ihren Text und fügen Sie Bilder ein"},
        {"ctx": "4.1 Kurze Einführung zum Interview", "md_re": r"^###?\s*(4\.1|Kurze\s+Einführung)", "ph": "Schreiben Sie hier Ihren Text…", "section_after": "Interview"},
        {"ctx": "4.2 Zusammenfassung des Interviews", "md_re": r"^###?\s*(4\.2|Zusammenfassung)", "ph": "Schreiben Sie hier Ihren Text…"},
        {"ctx": "4.3 Auswertung",                     "md_re": r"^###?\s*(4\.3|Auswertung)", "ph": "Schreiben Sie hier Ihren Text…"},
        {"ctx": "5.1 Umfrage Methode",                "md_re": r"^###?\s*(5\.1|Methode)", "ph": "Schreiben Sie hier Ihren Text…"},
        {"ctx": "5.2 Umfrage Ergebnisse",             "md_re": r"^###?\s*(5\.2|Ergebnisse|Rücklauf)", "ph": "Schreiben Sie hier Ihren Text…"},
        {"ctx": "5.3 Umfrage Interpretation",         "md_re": r"^###?\s*(5\.3|Interpretation)", "ph": "Schreiben Sie hier Ihren Text…"},
        {"ctx": "6.1 Schlusswort Zusammenfassung",    "md_re": r"^###?\s*(6\.1|Zusammenfassung)", "ph": "Schreiben Sie hier Ihren Text…", "section_after": "Schlusswort"},
        {"ctx": "6.2 Schlusswort Stellungnahme",      "md_re": r"^###?\s*(6\.2|Stellungnahme|Kommentar)", "ph": "Schreiben Sie hier Ihren Text…"},
        {"ctx": "Reflexion Arbeitsprozess",           "md_re": r"^##?\s*(Gesamtreflexion|Reflexion\s+des\s+Arbeitsprozesses)", "ph": "Schreiben Sie hier Ihren Text…", "section_after": "Reflexion des Arbeitsprozesses", "from_md": gesamtreflexion_md},
    ]

    last_found = 0
    for m in mapping:
        # Optional: search only after a specific heading
        search_from = last_found
        if "section_after" in m:
            h = _find_heading_idx(doc, m["section_after"], after=0)
            if h is not None:
                search_from = h

        # Source md
        src_md = m.get("from_md", haupttext_md)
        content = _extract_section(src_md, m["md_re"])
        if not content:
            log.warning(f"[docx] Keine MD-Section für {m['ctx']} gefunden (Regex: {m['md_re']})")
            continue

        blocks = _split_markdown_into_blocks(content)
        idx = _replace_placeholder_with_blocks(doc, m["ph"], blocks, after_idx=search_from)
        if idx is None:
            log.warning(f"[docx] Placeholder für {m['ctx']} nicht gefunden: '{m['ph']}'")
        else:
            last_found = idx + len(blocks)

    # === Quellenverzeichnis ===
    quellen_idx = _find_heading_idx(doc, "Quellenverzeichnis")
    if quellen_idx is not None:
        hint_idx = _find_para_idx_with_text(doc, "Orientieren Sie sich am gegebenen Beispiel", after=quellen_idx)
        if hint_idx is not None:
            # Alle Quellen als Liste
            lines = ["**Inhaltliche Quellen**", ""]
            by_kap: dict[str, list[str]] = {}
            for q in u.quellen:
                kap = q.kapitel_zuordnung or "Allgemein"
                by_kap.setdefault(kap, []).append(_format_source(q))
            for kap, srcs in by_kap.items():
                lines.append(f"*{kap}*")
                for s in srcs:
                    lines.append(f"- {s}")
                lines.append("")
            lines.append("**Personen**")
            lines.append(f"- Interview mit {u.interviewperson.name_anzeige}, {u.interviewperson.funktion}, {u.interviewperson.interview_termin}")
            lines.append(f"- Umfrage per {u.umfrage.plattform}, durchgeführt {u.umfrage.zeitraum}, {u.umfrage.n_ruecklauf} Teilnehmende")

            blocks = _split_markdown_into_blocks("\n".join(lines))
            _replace_placeholder_with_blocks(doc, "Orientieren Sie sich am gegebenen Beispiel", blocks, after_idx=quellen_idx)

    # === Konzept-Anhang ===
    konzept_anhang = _find_para_idx_with_text(doc, "Kopieren Sie Ihr eigenes unterschriebenes Konzept")
    if konzept_anhang is not None:
        blocks = _split_markdown_into_blocks(konzept_md)
        _replace_placeholder_with_blocks(doc, "Kopieren Sie Ihr eigenes unterschriebenes Konzept", blocks, after_idx=konzept_anhang)

    # === Tabelle 0: Name + Thema (im Zeitplan-Block) ===
    if len(doc.tables) >= 1:
        t = doc.tables[0]
        _fill_table_cell(t, 1, 0, f"{u.schuelerin.nachname}, {u.schuelerin.vorname}")
        _fill_table_cell(t, 1, 1, u.thema.titel)

    # === Tabelle 1: Zeitplan 22 Arbeitsschritte (nur Datum befüllen) ===
    if len(doc.tables) >= 2:
        t = doc.tables[1]
        # Zeilen 1..22 (erste Zeile ist Header). Datumpattern: pro Woche grob
        import datetime as dt
        start = dt.date(2025, 12, 1)
        for i in range(1, min(len(t.rows), 23)):
            due_date = start + dt.timedelta(weeks=(i - 1) // 3)
            done_date = due_date + dt.timedelta(days=3)
            _fill_table_cell(t, i, 2, due_date.strftime("%d.%m.%Y"))
            _fill_table_cell(t, i, 3, done_date.strftime("%d.%m.%Y"))

    # === Tabellen 2..9+: Wochenjournal (8 Wochen) ===
    # Parst aus journal_md 8 Wochen-Blöcke und befüllt je 1 Tabelle.
    journal_weeks = _parse_journal_weeks(journal_md, u)
    journal_tables = doc.tables[2:]  # ab Index 2 sind die Journal-Tabellen
    for week, table in zip(journal_weeks, journal_tables):
        _fill_journal_table(table, week)

    # === Zwischenreflexionen (Para 112-138 und 140-148 grob, je als Textblock direkt unter dem Heading) ===
    # Nach dem Heading "Erste Reflexion während des Schreibprozesses" einfügen
    z1_idx = _find_heading_idx(doc, "Erste Reflexion während des Schreibprozesses")
    if z1_idx is not None:
        # Erster leerer Heading-2 danach wird ersetzt
        next_empty = None
        for j in range(z1_idx + 1, min(z1_idx + 15, len(doc.paragraphs))):
            if doc.paragraphs[j].text.strip() == "" and doc.paragraphs[j].style.name.startswith("Heading"):
                next_empty = j
                break
        if next_empty is not None:
            blocks = _split_markdown_into_blocks(zwischenreflexion1_md)
            target = doc.paragraphs[next_empty]
            if blocks:
                _replace_paragraph_text(target, blocks[0]["text"])
                target.style = doc.styles["Normal"]
                prev = target
                for b in blocks[1:]:
                    prev = _insert_paragraph_after(prev, b["text"])

    z2_idx = _find_heading_idx(doc, "Zweite Reflexion während meines Schreibprozesses")
    if z2_idx is not None:
        next_empty = None
        for j in range(z2_idx + 1, min(z2_idx + 15, len(doc.paragraphs))):
            if doc.paragraphs[j].text.strip() == "" and doc.paragraphs[j].style.name.startswith("Heading"):
                next_empty = j
                break
        if next_empty is not None:
            blocks = _split_markdown_into_blocks(zwischenreflexion2_md)
            target = doc.paragraphs[next_empty]
            if blocks:
                _replace_paragraph_text(target, blocks[0]["text"])
                target.style = doc.styles["Normal"]
                prev = target
                for b in blocks[1:]:
                    prev = _insert_paragraph_after(prev, b["text"])

    # === Einverständniserklärung: Felder befüllen ===
    _replace_line_with_underscores(doc, "Thema der Vertiefungsarbeit:", u.thema.titel)
    _replace_line_with_underscores(doc, "Verfasser/in der Arbeit:", f"{u.schuelerin.vorname} {u.schuelerin.nachname}")
    _replace_line_with_underscores(doc, "Interviewte Person:", u.interviewperson.name_anzeige)
    _replace_line_with_underscores(doc, "Betreuende Lehrperson:", u.schuelerin.lehrperson)
    _replace_line_with_underscores(doc, "Datum des Interviews/der Aufnahme:", u.interviewperson.interview_termin)
    # Ort/Datum/Unterschrift Zeile
    _replace_line_with_underscores(doc, "……………………….", f"Winterthur\t\t{u.interviewperson.interview_termin}\t\t{u.interviewperson.name_anzeige}")

    # === Eigenleistungserklärung: Datum + Name ===
    eigen_idx = _find_heading_idx(doc, "Erklärung der Eigenleistung")
    if eigen_idx is not None:
        datum_idx = _find_para_idx_with_text(doc, "Datum", after=eigen_idx)
        if datum_idx is not None:
            _replace_paragraph_text(doc.paragraphs[datum_idx], "Datum: 15.04.2026")
        name_idx = _find_para_idx_with_text(doc, "Name und Unterschrift", after=eigen_idx)
        if name_idx is not None:
            _replace_paragraph_text(doc.paragraphs[name_idx], f"Name und Unterschrift: {u.schuelerin.vorname} {u.schuelerin.nachname}")

    # === Speichern ===
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    log.info(f"[docx] Haupt-VA geschrieben: {out_path}")
    return out_path


def _format_source(q) -> str:
    """APA-light-Formatierung einer Quelle."""
    if q.typ == "buch":
        parts = [q.autor, f"({q.jahr})" if q.jahr else "", q.titel]
        if q.verlag:
            parts.append(q.verlag)
        if q.isbn:
            parts.append(f"ISBN {q.isbn}")
        return " ".join(p for p in parts if p).replace("  ", " ").strip(". ") + "."
    if q.typ == "fachartikel":
        parts = [q.autor, f"({q.jahr})" if q.jahr else "", q.titel]
        if q.verlag:
            parts.append(f"In: {q.verlag}")
        if q.doi:
            parts.append(f"DOI: {q.doi}")
        elif q.url:
            parts.append(q.url)
        return " ".join(p for p in parts if p).strip(". ") + "."
    if q.typ == "internet":
        parts = [q.autor, f"({q.jahr})" if q.jahr else "", q.titel]
        if q.url:
            parts.append(q.url)
        return " ".join(p for p in parts if p).strip(". ") + "."
    return f"{q.autor}: {q.titel} ({q.jahr or 'o. J.'}). {q.url or q.verlag or ''}".strip(". ") + "."


def _parse_journal_weeks(journal_md: str, u: Universe) -> list[dict]:
    """Parst 8 Wochenblöcke aus dem Journal-Markdown."""
    weeks: list[dict] = []
    # Teile am "## Woche N" — flexibel
    chunks = re.split(r"^##\s*Woche\s+\d", journal_md, flags=re.MULTILINE)
    headers = re.findall(r"^##\s*Woche\s+\d+[^\n]*", journal_md, flags=re.MULTILINE)
    for i, header in enumerate(headers):
        # Chunks[0] ist der Einleitungstext vor der ersten Woche
        body = chunks[i + 1] if i + 1 < len(chunks) else ""
        datum = ""
        # Versuche Datum aus dem Header zu ziehen: "Woche 1 · 2025-12-01"
        date_match = re.search(r"(\d{1,2}[.\-/]\d{1,2}[.\-/]\d{2,4}|\d{4}-\d{2}-\d{2})", header)
        if date_match:
            datum = date_match.group(1)
        elif i < len(u.timeline):
            datum = u.timeline[i].datum_start

        # Extrahiere die drei Sub-Sections
        taetigkeiten = _extract_subsection(body, r"Tätigkeiten")
        zu_erledigen = _extract_subsection(body, r"Zu erledigen")
        naechster_unterricht = _extract_subsection(body, r"nächsten Unterricht")

        weeks.append({
            "datum": datum,
            "taetigkeiten": taetigkeiten or body[:500],
            "zu_erledigen": zu_erledigen,
            "naechster_unterricht": naechster_unterricht,
        })
    return weeks[:12]  # Max 12


def _extract_subsection(body: str, heading_re: str) -> str:
    """Extrahiert eine Unter-Section aus einem Wochenblock."""
    lines = body.splitlines()
    in_section = False
    collected: list[str] = []
    for line in lines:
        if re.search(heading_re, line, flags=re.IGNORECASE) and (line.startswith("###") or "**" in line):
            in_section = True
            continue
        if in_section and (line.startswith("###") or line.startswith("##")):
            break
        if in_section:
            collected.append(line)
    return "\n".join(collected).strip()


def _fill_journal_table(table, week: dict) -> None:
    """Füllt eine 6×2 Wochenjournal-Tabelle.

    Struktur (nach Vorlage):
    Zeile 0: DATUM | "Tätigkeiten / Dauer / Vorgehen"
    Zeile 1: DATUM | <Text Tätigkeiten>
    Zeile 2: DATUM | "Zu erledigen bis nächste Woche"
    Zeile 3: DATUM | <Text Zu erledigen>
    Zeile 4: DATUM | "Arbeiten für den nächsten Unterricht"
    Zeile 5: DATUM | <Text>
    """
    datum = week["datum"]
    # Alle "DATUM"-Zellen auf echtes Datum setzen
    for r in range(len(table.rows)):
        cell0 = table.rows[r].cells[0]
        if cell0.text.strip() == "DATUM":
            cell0.text = datum
    # Inhaltszellen
    try:
        if len(table.rows) >= 2:
            table.rows[1].cells[1].text = week.get("taetigkeiten", "")
        if len(table.rows) >= 4:
            table.rows[3].cells[1].text = week.get("zu_erledigen", "")
        if len(table.rows) >= 6:
            table.rows[5].cells[1].text = week.get("naechster_unterricht", "")
    except Exception as e:
        log.warning(f"Journal-Tabelle konnte nicht vollständig gefüllt werden: {e}")


def _replace_line_with_underscores(doc: Document, prefix: str, replacement: str) -> None:
    """Findet einen Paragraph der mit `prefix` beginnt und ersetzt die '____' durch `replacement`."""
    idx = _find_para_idx_with_text(doc, prefix)
    if idx is None:
        return
    p = doc.paragraphs[idx]
    new_text = re.sub(r"_{5,}", replacement, p.text, count=1)
    if new_text == p.text:
        new_text = re.sub(r"…+", replacement, p.text, count=1)
    _replace_paragraph_text(p, new_text)


def docx_to_pdf(docx_path: Path) -> Path | None:
    """Versucht .docx → .pdf via LibreOffice. Liefert None wenn kein soffice verfügbar."""
    pdf_path = docx_path.with_suffix(".pdf")
    for binary in ("/Applications/LibreOffice.app/Contents/MacOS/soffice", "soffice", "libreoffice"):
        try:
            result = subprocess.run(
                [binary, "--headless", "--convert-to", "pdf", "--outdir", str(pdf_path.parent), str(docx_path)],
                capture_output=True, timeout=90,
            )
            if result.returncode == 0 and pdf_path.exists():
                log.info(f"[docx] → PDF erfolgreich: {pdf_path}")
                return pdf_path
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue
    log.warning("[docx] LibreOffice nicht gefunden, kein PDF erstellt")
    return None
